import argparse
import json
from pathlib import Path


def read_pdf(path: Path):
    try:
        from pypdf import PdfReader
    except ImportError:
        return {"error": "Missing dependency: pypdf. Run pip install -r requirements.txt"}
    reader = PdfReader(str(path))
    pages = [(page.extract_text() or "").strip() for page in reader.pages]
    return {"type": "pdf", "pages": len(pages), "text": "\n\n".join(pages)}


def read_docx(path: Path):
    try:
        from docx import Document
    except ImportError:
        return {"error": "Missing dependency: python-docx. Run pip install -r requirements.txt"}
    doc = Document(str(path))
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return {"type": "docx", "paragraphs": len(doc.paragraphs), "text": text}


def read_text(path: Path):
    return {"type": path.suffix.lower().lstrip(".") or "text", "text": path.read_text(encoding="utf-8", errors="replace")}


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("path")
    parser.add_argument("--max-chars", type=int, default=12000)
    args = parser.parse_args()

    path = Path(args.path.lstrip("/"))
    if not path.exists():
        data = {"error": f"File not found: {path}"}
    elif path.suffix.lower() == ".pdf":
        data = read_pdf(path)
    elif path.suffix.lower() == ".docx":
        data = read_docx(path)
    else:
        data = read_text(path)

    if "text" in data and len(data["text"]) > args.max_chars:
        data["text"] = data["text"][: args.max_chars] + "\n... [truncated]"
    print(json.dumps(data, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
